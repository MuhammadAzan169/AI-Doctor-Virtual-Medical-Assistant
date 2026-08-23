"""Lab-report text extraction for PDF and DOCX uploads.

Images go straight to PaddleOCR. Documents are handled here first, because
both PDF and DOCX usually carry a real text layer — reading that is exact and
takes milliseconds, where OCR is lossy and takes tens of seconds. A PDF that
turns out to be a scan (no meaningful text layer) falls back to rasterising
each page and running OCR over it.
"""

import os

from app.core.logging import get_logger
from app.services import ocr

logger = get_logger("AIDoctor.Documents")

# A PDF whose text layer yields less than this is treated as a scan.
_TEXT_LAYER_MIN_CHARS = 120

# Rasterisation DPI for scanned PDFs. 150 keeps small print legible while
# staying under the OCR input cap for A4.
_PDF_RASTER_DPI = 150

# Never OCR an unbounded page count; reports are short and each page is slow.
_MAX_PDF_PAGES = 10


def extract_docx_text(path: str) -> str:
    """Read paragraphs and tables out of a .docx in document order.

    Walks the body XML rather than doc.paragraphs + doc.tables, which would
    return every paragraph first and every table afterwards — that separates
    each panel heading from the results underneath it.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    lines = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = Paragraph(child, doc).text.strip()
            if text:
                lines.append(text)
        elif tag == "tbl":
            for row in Table(child, doc).rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append("  ".join(cells))
    return chr(10).join(lines)


def _pdf_text_layer(path: str) -> str:
    """Return embedded PDF text, or '' when there is none worth using."""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(path)
    try:
        pages = []
        for i in range(min(len(pdf), _MAX_PDF_PAGES)):
            textpage = pdf[i].get_textpage()
            try:
                pages.append(textpage.get_text_bounded())
            finally:
                textpage.close()
        return "\n".join(p for p in pages if p and p.strip()).strip()
    finally:
        pdf.close()


def _pdf_page_images(path: str, output_dir: str) -> list[str]:
    """Rasterise PDF pages to PNGs and return their paths."""
    import pypdfium2 as pdfium

    os.makedirs(output_dir, exist_ok=True)
    stem = os.path.splitext(os.path.basename(path))[0]
    pdf = pdfium.PdfDocument(path)
    try:
        paths = []
        for i in range(min(len(pdf), _MAX_PDF_PAGES)):
            image = pdf[i].render(scale=_PDF_RASTER_DPI / 72).to_pil()
            page_path = os.path.join(output_dir, f"{stem}_page{i + 1}.png")
            image.save(page_path)
            paths.append(page_path)
        return paths
    finally:
        pdf.close()


def extract_pdf_text(path: str, output_dir: str) -> str:
    """Prefer the PDF's own text layer; OCR the rendered pages when it is a scan."""
    try:
        text = _pdf_text_layer(path)
    except Exception:
        logger.exception("Reading the PDF text layer failed for %s", path)
        text = ""

    if len(text) >= _TEXT_LAYER_MIN_CHARS:
        logger.info("PDF text layer used (%d chars) for %s", len(text), path)
        return text

    if not ocr.is_available():
        logger.warning("PDF %s has no usable text layer and OCR is disabled", path)
        return (
            "This PDF appears to be a scan and text extraction (OCR) is not enabled "
            "on this server, so its contents could not be read."
        )

    logger.info("PDF text layer too thin (%d chars) — rasterising and running OCR", len(text))
    try:
        page_images = _pdf_page_images(path, output_dir)
    except Exception:
        logger.exception("Rasterising %s failed", path)
        return "[PDF processing failed]"

    pages = []
    for page_path in page_images:
        json_path = ocr.perform_ocr(page_path, output_dir)
        if json_path and os.path.exists(json_path):
            pages.append(ocr.extract_text_from_json(json_path))
    return "\n".join(p for p in pages if p) or "[OCR processing failed]"


def extract_document_text(path: str, output_dir: str) -> str:
    """Dispatch a .pdf or .docx upload to the right extractor."""
    suffix = os.path.splitext(path)[1].lower()
    if suffix == ".docx":
        try:
            return extract_docx_text(path) or "[The document contained no readable text]"
        except Exception:
            logger.exception("DOCX extraction failed for %s", path)
            return "[DOCX processing failed]"
    if suffix == ".pdf":
        return extract_pdf_text(path, output_dir)
    raise ValueError(f"extract_document_text does not handle {suffix}")
