# -*- coding: utf-8 -*-
"""Generate a sample lab report (PDF + PNG) for testing AI Doctor uploads."""
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.pdfgen import canvas

W, H = A4
OUT = "sample_data/sample_lab_report.pdf"

NAVY = colors.HexColor("#123a5e")
GREY = colors.HexColor("#555555")
LINE = colors.HexColor("#c9d3dd")
RED  = colors.HexColor("#b32020")

PANELS = [
    ("COMPLETE BLOOD COUNT (CBC)", [
        ("Haemoglobin",            "11.2",  "g/dL",     "13.0 - 17.0", "LOW"),
        ("Total Leucocyte Count",  "12,800","/uL",      "4,000 - 11,000", "HIGH"),
        ("Neutrophils",            "78",    "%",        "40 - 75",     "HIGH"),
        ("Lymphocytes",            "16",    "%",        "20 - 40",     "LOW"),
        ("Eosinophils",            "2",     "%",        "1 - 6",       ""),
        ("Monocytes",              "4",     "%",        "2 - 10",      ""),
        ("Platelet Count",         "1.9",   "lakh/uL",  "1.5 - 4.1",   ""),
        ("RBC Count",              "4.1",   "mil/uL",   "4.5 - 5.5",   "LOW"),
        ("MCV",                    "74.5",  "fL",       "83 - 101",    "LOW"),
        ("MCH",                    "24.1",  "pg",       "27 - 32",     "LOW"),
    ]),
    ("BASIC METABOLIC PANEL", [
        ("Fasting Blood Glucose",  "126",   "mg/dL",    "70 - 100",    "HIGH"),
        ("HbA1c",                  "6.8",   "%",        "< 5.7",       "HIGH"),
        ("Blood Urea",             "28",    "mg/dL",    "17 - 43",     ""),
        ("Serum Creatinine",       "1.0",   "mg/dL",    "0.7 - 1.3",   ""),
        ("Sodium",                 "138",   "mmol/L",   "136 - 145",   ""),
        ("Potassium",              "4.2",   "mmol/L",   "3.5 - 5.1",   ""),
    ]),
    ("LIPID PROFILE", [
        ("Total Cholesterol",      "214",   "mg/dL",    "< 200",       "HIGH"),
        ("HDL Cholesterol",        "38",    "mg/dL",    "> 40",        "LOW"),
        ("LDL Cholesterol",        "142",   "mg/dL",    "< 100",       "HIGH"),
        ("Triglycerides",          "168",   "mg/dL",    "< 150",       "HIGH"),
    ]),
    ("INFLAMMATORY MARKERS", [
        ("C-Reactive Protein",     "18.4",  "mg/L",     "< 5.0",       "HIGH"),
        ("ESR",                    "34",    "mm/hr",    "0 - 15",      "HIGH"),
    ]),
]

c = canvas.Canvas(OUT, pagesize=A4)
c.setTitle("Sample Lab Report")
c.setAuthor("AI Doctor - test fixture")
c.setSubject("Synthetic laboratory report for application testing")

# ---- header band -------------------------------------------------------
def watermark():
    c.saveState()
    c.setFillColor(colors.HexColor("#123a5e"), alpha=0.06)
    c.setFont("Helvetica-Bold", 78)
    c.translate(W/2, H/2); c.rotate(38)
    c.drawCentredString(0, 0, "SAMPLE")
    c.restoreState()

watermark()
c.setFillColor(NAVY)
c.rect(0, H - 30*mm, W, 30*mm, stroke=0, fill=1)
c.setFillColor(colors.white)
c.setFont("Helvetica-Bold", 16)
c.drawString(18*mm, H - 14*mm, "MERIDIAN DIAGNOSTIC LABORATORY")
c.setFont("Helvetica", 8.5)
c.drawString(18*mm, H - 20*mm, "42 Clinic Road, Springfield  |  +1 555 0100  |  reports@meridian-labs.example")
c.setStrokeColor(colors.white); c.setLineWidth(0.9)
c.rect(146*mm, H - 17.5*mm, 46*mm, 8*mm, stroke=1, fill=0)
c.setFont("Helvetica-Bold", 8.5)
c.drawCentredString(169*mm, H - 15*mm, "SAMPLE - NOT A REAL REPORT")
c.setFont("Helvetica", 7.5)
c.drawCentredString(169*mm, H - 22*mm, "Fictitious data - software testing only")

# ---- patient block -----------------------------------------------------
y = H - 40*mm
def field(x, y, label, value):
    c.setFillColor(GREY); c.setFont("Helvetica", 8.5)
    c.drawString(x, y, label)
    c.setFillColor(colors.black); c.setFont("Helvetica-Bold", 10)
    c.drawString(x, y - 5*mm, value)

field(18*mm,  y, "PATIENT NAME",  "John Q. Sample")
field(78*mm,  y, "AGE / SEX",     "34 Y / Male")
field(120*mm, y, "PATIENT ID",    "MDL-2026-004178")
field(160*mm, y, "SAMPLE TYPE",   "Whole Blood")
y -= 13*mm
field(18*mm,  y, "REFERRED BY",   "Dr. A. Placeholder, MBBS")
field(78*mm,  y, "COLLECTED ON",  "24 Aug 2026, 08:15")
field(120*mm, y, "REPORTED ON",   "24 Aug 2026, 14:40")
field(160*mm, y, "ACCESSION",     "AC-77120934")

y -= 10*mm
c.setStrokeColor(LINE); c.setLineWidth(0.8)
c.line(18*mm, y, W - 18*mm, y)
y -= 9*mm

COLS = (18*mm, 92*mm, 118*mm, 140*mm, 178*mm)  # test, result, unit, range, flag

def header_row(yy):
    c.setFillColor(colors.HexColor("#eef3f8"))
    c.rect(18*mm, yy - 2.5*mm, W - 36*mm, 7*mm, stroke=0, fill=1)
    c.setFillColor(NAVY); c.setFont("Helvetica-Bold", 8.5)
    c.drawString(COLS[0] + 2*mm, yy, "TEST")
    c.drawRightString(COLS[1], yy, "RESULT")
    c.drawString(COLS[2], yy, "UNIT")
    c.drawString(COLS[3], yy, "REFERENCE RANGE")
    c.drawString(COLS[4], yy, "FLAG")
    return yy - 6.4*mm

for title, rows in PANELS:
    if y < 45*mm:
        c.showPage(); watermark(); y = H - 25*mm
    c.setFillColor(NAVY); c.setFont("Helvetica-Bold", 10.5)
    c.drawString(18*mm, y, title)
    y -= 5.5*mm
    y = header_row(y)
    for i, (name, res, unit, ref, flag) in enumerate(rows):
        if i % 2 == 0:
            c.setFillColor(colors.HexColor("#fafbfc"))
            c.rect(18*mm, y - 1.7*mm, W - 36*mm, 5.2*mm, stroke=0, fill=1)
        c.setFillColor(colors.black); c.setFont("Helvetica", 9)
        c.drawString(COLS[0] + 2*mm, y, name)
        c.setFont("Helvetica-Bold" if flag else "Helvetica", 9)
        c.setFillColor(RED if flag else colors.black)
        c.drawRightString(COLS[1], y, res)
        c.setFillColor(GREY); c.setFont("Helvetica", 8.5)
        c.drawString(COLS[2], y, unit)
        c.drawString(COLS[3], y, ref)
        if flag:
            c.setFillColor(RED); c.setFont("Helvetica-Bold", 8.5)
            c.drawString(COLS[4], y, flag)
        y -= 5.2*mm
    y -= 3*mm

# ---- interpretation ----------------------------------------------------
if y < 46*mm:
    c.showPage(); watermark(); y = H - 25*mm
c.setFillColor(NAVY); c.setFont("Helvetica-Bold", 10.5)
c.drawString(18*mm, y, "INTERPRETATION")
y -= 6*mm
c.setFillColor(colors.black); c.setFont("Helvetica", 9)
for line in [
    "Microcytic hypochromic anaemia pattern with low haemoglobin, MCV and MCH - consider iron studies.",
    "Neutrophilic leucocytosis with raised CRP and ESR suggests an active bacterial infection.",
    "Fasting glucose and HbA1c are in the diabetic range; correlate clinically and repeat to confirm.",
    "Dyslipidaemia noted: raised LDL and triglycerides with low HDL.",
]:
    c.drawString(20*mm, y, u"\u2022  " + line)
    y -= 5*mm

y -= 4*mm

# signature sits above the rule, on its own band
c.setFont("Helvetica-Bold", 9); c.setFillColor(colors.black)
c.drawRightString(W - 18*mm, y, "Dr. A. Placeholder, MD (Pathology)")
c.setFont("Helvetica", 8); c.setFillColor(GREY)
c.drawRightString(W - 18*mm, y - 4.5*mm, "Consultant Pathologist  |  Reg. No. SAMPLE-0000")

y -= 11*mm
c.setStrokeColor(LINE); c.setLineWidth(0.8)
c.line(18*mm, y, W - 18*mm, y)
y -= 5*mm
c.setFillColor(GREY); c.setFont("Helvetica-Oblique", 7.5)
for line in [
    "This document is synthetic test data generated for the AI Doctor project. The patient, physician and laboratory",
    "named above do not exist, and the results describe no real person. Not for clinical or diagnostic use.",
]:
    c.drawString(18*mm, y, line)
    y -= 4*mm

c.save()
print("wrote", OUT)

# ---- DOCX variant of the same report -----------------------------------
# Same data, different container, so the .docx upload path can be tested
# against a known-good expected text.
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.core_properties.title = "Sample Lab Report"
doc.core_properties.comments = "Synthetic test data for the AI Doctor project."

h = doc.add_paragraph()
run = h.add_run("MERIDIAN DIAGNOSTIC LABORATORY")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x12, 0x3A, 0x5E)

sub = doc.add_paragraph()
sub_run = sub.add_run("SAMPLE - NOT A REAL REPORT  |  Fictitious data, for software testing only")
sub_run.bold = True
sub_run.font.size = Pt(9)
sub_run.font.color.rgb = RGBColor(0xB3, 0x20, 0x20)

info = doc.add_table(rows=2, cols=4)
info.style = "Table Grid"
for cell, text in zip(info.rows[0].cells, [
    "Patient: John Q. Sample", "Age / Sex: 34 Y / Male",
    "Patient ID: MDL-2026-004178", "Sample: Whole Blood"]):
    cell.text = text
for cell, text in zip(info.rows[1].cells, [
    "Referred by: Dr. A. Placeholder, MBBS", "Collected: 24 Aug 2026, 08:15",
    "Reported: 24 Aug 2026, 14:40", "Accession: AC-77120934"]):
    cell.text = text

for title, rows in PANELS:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["TEST", "RESULT", "UNIT", "REFERENCE RANGE", "FLAG"]):
        cell.paragraphs[0].add_run(text).bold = True
    for name, res, unit, ref, flag in rows:
        cells = table.add_row().cells
        cells[0].text = name
        run = cells[1].paragraphs[0].add_run(res)
        if flag:
            run.bold = True
            run.font.color.rgb = RGBColor(0xB3, 0x20, 0x20)
        cells[2].text = unit
        cells[3].text = ref
        if flag:
            flag_run = cells[4].paragraphs[0].add_run(flag)
            flag_run.bold = True
            flag_run.font.color.rgb = RGBColor(0xB3, 0x20, 0x20)

doc.add_heading("INTERPRETATION", level=2)
for line in [
    "Microcytic hypochromic anaemia pattern with low haemoglobin, MCV and MCH - consider iron studies.",
    "Neutrophilic leucocytosis with raised CRP and ESR suggests an active bacterial infection.",
    "Fasting glucose and HbA1c are in the diabetic range; correlate clinically and repeat to confirm.",
    "Dyslipidaemia noted: raised LDL and triglycerides with low HDL.",
]:
    doc.add_paragraph(line, style="List Bullet")

sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig.add_run("Dr. A. Placeholder, MD (Pathology)\nConsultant Pathologist  |  Reg. No. SAMPLE-0000").bold = True

foot = doc.add_paragraph()
foot_run = foot.add_run(
    "This document is synthetic test data generated for the AI Doctor project. The patient, physician "
    "and laboratory named above do not exist, and the results describe no real person. "
    "Not for clinical or diagnostic use."
)
foot_run.italic = True
foot_run.font.size = Pt(7.5)

DOCX_OUT = "sample_data/sample_lab_report.docx"
doc.save(DOCX_OUT)
print("wrote", DOCX_OUT)
